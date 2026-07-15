"""
DOCX text extractor — extracts text from DOCX files using python-docx.

Returns concatenated paragraph text preserving document order.
"""

import logging
from pathlib import Path
from typing import Union

from app.exceptions.base import DocumentProcessingException

logger = logging.getLogger(__name__)


def extract_text_from_docx(
    source: Union[str, Path, bytes],
) -> str:
    """
    Extract all paragraph text from a DOCX file.

    Parameters
    ----------
    source : str | Path | bytes
        File path or raw byte content.

    Returns
    -------
    str
        Concatenated paragraph text.

    Raises
    ------
    DocumentProcessingException
        If the file cannot be read.
    """
    try:
        import io

        from docx import Document as DocxDocument
    except ImportError as exc:
        raise DocumentProcessingException(message="python-docx is not installed; cannot extract DOCX text.") from exc

    try:
        if isinstance(source, bytes):
            doc = DocxDocument(io.BytesIO(source))
        else:
            doc = DocxDocument(str(source))

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        if not paragraphs:
            raise DocumentProcessingException(message="DOCX file contains no readable text content.")

        full_text = "\n\n".join(paragraphs)
        logger.info(
            "DOCX text extraction complete: %d characters across %d paragraphs",
            len(full_text),
            len(paragraphs),
        )
        return full_text

    except DocumentProcessingException:
        raise
    except Exception as exc:
        raise DocumentProcessingException(message=f"Failed to extract text from DOCX file: {exc}") from exc
