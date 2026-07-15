"""
Table extractor — extracts tabular data from PDF and DOCX documents.

Returns tables as ``List[List[List[str]]]`` — a list of tables, each
table being a list of rows, each row being a list of cell strings.
"""

import io
import logging
from pathlib import Path
from typing import List, Union

logger = logging.getLogger(__name__)

# Type alias: a single table is rows × cells
Table = List[List[str]]


def extract_tables_from_pdf(
    source: Union[str, Path, bytes],
) -> List[Table]:
    """
    Extract all tables from a PDF using pdfplumber.

    Parameters
    ----------
    source : str | Path | bytes
        File path or raw PDF bytes.

    Returns
    -------
    list[Table]
        Each Table is a list of rows (list of cell strings).
    """
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber is not installed; cannot extract PDF tables.")
        return []

    try:
        if isinstance(source, bytes):
            pdf = pdfplumber.open(io.BytesIO(source))
        else:
            pdf = pdfplumber.open(str(source))

        all_tables: List[Table] = []

        for page_num, page in enumerate(pdf.pages, start=1):
            page_tables = page.extract_tables() or []
            for tbl_idx, raw_table in enumerate(page_tables):
                if not raw_table:
                    continue
                cleaned_table: Table = []
                for row in raw_table:
                    cleaned_row = [(cell.strip() if cell else "") for cell in row]
                    cleaned_table.append(cleaned_row)
                all_tables.append(cleaned_table)

        pdf.close()
        logger.info("Extracted %d tables from PDF", len(all_tables))
        return all_tables

    except Exception as exc:
        logger.warning("PDF table extraction failed: %s", exc)
        return []


def extract_tables_from_docx(
    source: Union[str, Path, bytes],
) -> List[Table]:
    """
    Extract all tables from a DOCX file using python-docx.

    Parameters
    ----------
    source : str | Path | bytes
        File path or raw DOCX bytes.

    Returns
    -------
    list[Table]
        Each Table is a list of rows (list of cell strings).
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.warning("python-docx is not installed; cannot extract DOCX tables.")
        return []

    try:
        if isinstance(source, bytes):
            doc = DocxDocument(io.BytesIO(source))
        else:
            doc = DocxDocument(str(source))

        all_tables: List[Table] = []

        for tbl in doc.tables:
            cleaned_table: Table = []
            for row in tbl.rows:
                cleaned_row = [cell.text.strip() for cell in row.cells]
                cleaned_table.append(cleaned_row)
            if cleaned_table:
                all_tables.append(cleaned_table)

        logger.info("Extracted %d tables from DOCX", len(all_tables))
        return all_tables

    except Exception as exc:
        logger.warning("DOCX table extraction failed: %s", exc)
        return []
